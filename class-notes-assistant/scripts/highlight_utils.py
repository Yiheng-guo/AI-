"""
高亮颜色工具函数
用于生成带颜色高亮的 Word 文档笔记

颜色规范（用户锁定）：
- 🟡 黄色 (highlight_color=7) = 特别重点/核心理念/面试话术，加粗
- 🟢 绿色 (highlight_color=4) = 重点问题，加粗，极少用只标真正关键的
- 🔵 蓝色 (highlight_color=2) = 重点问题的回答
- 🟣 紫色 (highlight_color=12) = 面试答题要点/tips，加粗
"""

from docx import Document
from docx.shared import Pt


class HighlightWriter:
    """带高亮的文档写入器"""

    # 颜色常量
    YELLOW = 7   # 特别重点
    GREEN = 4    # 重点问题
    BLUE = 2     # 问题回答
    PURPLE = 12  # 答题要点

    def __init__(self, font_name='微软雅黑', font_size=11):
        self.doc = Document()
        self.font_name = font_name
        self.font_size = font_size

        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)

    def _set_run_style(self, run, bold=False, color=None):
        """设置 run 的字体样式"""
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        if bold:
            run.bold = True
        if color is not None:
            run.font.highlight_color = color

    def yellow(self, text):
        """🟡 黄色高亮：特别重点/核心理念，加粗"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_run_style(run, bold=True, color=self.YELLOW)
        return p

    def green(self, text):
        """🟢 绿色高亮：重点问题，加粗"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_run_style(run, bold=True, color=self.GREEN)
        return p

    def blue(self, text):
        """🔵 蓝色高亮：重点问题的回答"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_run_style(run, color=self.BLUE)
        return p

    def purple(self, text):
        """🟣 紫色高亮：面试答题要点/tips，加粗"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_run_style(run, bold=True, color=self.PURPLE)
        return p

    def qa_block(self, question, answer, tips=None):
        """
        标准问答块（用户确认的新格式）
        顺序：紫色答题要点 → 绿色问题 → 蓝色完整答案

        Args:
            question: 问题文本
            answer: 完整答案（自然流畅的段落，不用引号包裹）
            tips: 答题要点/小技巧（可选）
        """
        if tips:
            self.purple(f'💡 答题要点：{tips}')
        self.green(question)
        self.blue(answer)
        self.doc.add_paragraph()  # 空行分隔

    def h1(self, text):
        """一级标题"""
        p = self.doc.add_heading(text, level=1)
        for run in p.runs:
            run.font.name = self.font_name
        return p

    def h2(self, text):
        """二级标题"""
        p = self.doc.add_heading(text, level=2)
        for run in p.runs:
            run.font.name = self.font_name
        return p

    def h3(self, text):
        """三级标题"""
        p = self.doc.add_heading(text, level=3)
        for run in p.runs:
            run.font.name = self.font_name
        return p

    def p(self, text):
        """普通段落"""
        para = self.doc.add_paragraph(text)
        for run in para.runs:
            run.font.name = self.font_name
        return para

    def bullet(self, text):
        """项目符号"""
        para = self.doc.add_paragraph(text, style='List Bullet')
        for run in para.runs:
            run.font.name = self.font_name
        return para

    def add_table(self, headers, rows):
        """
        添加表格

        Args:
            headers: 表头列表
            rows: 数据行列表（每行是一个列表）
        """
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'

        # 表头
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.name = self.font_name

        # 数据行
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = val
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = self.font_name

        return table

    def save(self, path):
        """保存文档"""
        self.doc.save(path)
        print(f'笔记已保存到：{path}')
